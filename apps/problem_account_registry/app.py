from __future__ import annotations

import hashlib
import html
import json
import os
import re
import sys
import uuid
from datetime import datetime
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import quote, unquote, urlparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import Workbook, load_workbook
from openpyxl.comments import Comment
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.worksheet.table import Table, TableStyleInfo


ROOT = Path(os.environ.get("K_DESK_ROOT", Path(__file__).resolve().parents[2]))
OUT_DIR = Path(os.environ.get("ACCOUNT_REGISTRY_DATA_DIR", ROOT / "local_data" / "problem_account_registry"))
KLINE_OUT_DIR = Path(os.environ.get("TRADE_KLINE_OUT_DIR", ROOT / "outputs" / "kline"))
TRADE_KLINE_WEB_URL = os.environ.get("TRADE_KLINE_WEB_URL", "http://127.0.0.1:8765")
SOURCE_TXT = Path(os.environ.get("ACCOUNT_REGISTRY_SOURCE_TXT", ROOT / "local_data" / "source_notes.txt"))
WORKBOOK_PATH = OUT_DIR / "problematic_accounts.xlsx"
HOST = "127.0.0.1"
PORT = int(os.environ.get("ACCOUNT_REGISTRY_PORT", "8776"))

SHEET_ACCOUNTS = "问题账户"
SHEET_RAW = "原始记录"
SHEET_HELP = "字段说明"
SHEET_HISTORY = "修改历史"

HEADERS = [
    "记录ID",
    "账号",
    "记录类型",
    "关联账号/主体",
    "建议动作",
    "当前分组",
    "风险标签",
    "风险/问题备注",
    "原始记录",
    "加入时间",
    "修改时间",
    "状态",
    "处理人/来源",
]

LEGACY_HEADER_ALIASES = {
    "加入时间": ["加入时间", "首次录入日期"],
    "修改时间": ["修改时间", "更新时间"],
}

HISTORY_HEADERS = [
    "历史ID",
    "记录ID",
    "账号",
    "操作",
    "修改时间",
    "修改字段",
    "修改前JSON",
    "修改后JSON",
    "处理人/来源",
]

EDITABLE_FIELDS = [
    "账号",
    "记录类型",
    "关联账号/主体",
    "建议动作",
    "当前分组",
    "风险标签",
    "风险/问题备注",
    "原始记录",
    "状态",
    "处理人/来源",
]

ACTION_CHOICES = [
    "",
    "M",
    "M观察",
    "P",
    "P观察",
    "T",
    "B-M",
    "B-P",
    "M-P",
    "P->A/T",
    "限制出金",
    "自定义",
    "待定",
]
STATUS_CHOICES = ["待复核", "观察中", "已确认", "已关闭"]
TYPE_CHOICES = ["账户", "IB/组", "其他"]
JOURNAL_DIR = Path(r"C:\Users\amber\Downloads")


def now_text() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def normalize_datetime(value: object, default_time: str = "12:00:00") -> str:
    text = normalize_text(value)
    if not text:
        return ""
    text = text.replace("/", "-")
    if re.fullmatch(r"\d{8}", text):
        return f"{text[:4]}-{text[4:6]}-{text[6:]} {default_time}"
    if re.fullmatch(r"\d{4}-\d{1,2}-\d{1,2}", text):
        dt = datetime.strptime(text, "%Y-%m-%d")
        return dt.strftime(f"%Y-%m-%d {default_time}")
    if re.fullmatch(r"\d{4}-\d{1,2}-\d{1,2}\s+\d{1,2}:\d{1,2}", text):
        dt = datetime.strptime(text, "%Y-%m-%d %H:%M")
        return dt.strftime("%Y-%m-%d %H:%M:%S")
    if re.fullmatch(r"\d{4}-\d{1,2}-\d{1,2}\s+\d{1,2}:\d{1,2}:\d{1,2}", text):
        dt = datetime.strptime(text, "%Y-%m-%d %H:%M:%S")
        return dt.strftime("%Y-%m-%d %H:%M:%S")
    return text


def account_join_times_from_journals() -> dict[str, str]:
    mapping: dict[str, str] = {}
    group_mapping: dict[str, str] = {}
    for path in sorted(JOURNAL_DIR.glob("journal_202606*.md")):
        match = re.search(r"(\d{8})", path.name)
        if not match:
            continue
        joined_at = normalize_datetime(match.group(1))
        text = path.read_text(encoding="utf-8", errors="ignore")
        for line in text.splitlines():
            if line.startswith("[image") or "base64," in line:
                continue
            if line.strip().lower().startswith("ib "):
                group_mapping.setdefault("IB/组", joined_at)
            for account_id in re.findall(r"\d{5,10}", line):
                mapping.setdefault(account_id, joined_at)

    # These source rows say "同上" in the journal and do not repeat every account id.
    for account_id in ("241002225", "5002797", "634847"):
        mapping.setdefault(account_id, "2026-06-22 12:00:00")
    mapping["5006543"] = "2026-06-23 10:30:00"
    mapping.update({f"__GROUP__{key}": value for key, value in group_mapping.items()})
    return mapping


def default_join_time(record: dict[str, str], journal_times: dict[str, str] | None = None) -> str:
    journal_times = journal_times or account_join_times_from_journals()
    account_id = normalize_text(record.get("账号"))
    if account_id and account_id in journal_times:
        return journal_times[account_id]
    if record.get("记录类型") == "IB/组" and "__GROUP__IB/组" in journal_times:
        return journal_times["__GROUP__IB/组"]
    existing = normalize_text(record.get("加入时间") or record.get("首次录入日期"))
    if existing and existing != "2026-06-23 12:00:00":
        return normalize_datetime(existing)
    raw = normalize_text(record.get("原始记录"))
    for account_id in re.findall(r"\b\d{5,10}\b", raw):
        if account_id in journal_times:
            return journal_times[account_id]
    return "2026-06-16 12:00:00"


def read_source_text() -> str:
    data = SOURCE_TXT.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def normalize_text(value: object) -> str:
    return str(value or "").strip()


def uniq_join(values: list[str], sep: str = "；") -> str:
    seen: set[str] = set()
    out: list[str] = []
    for value in values:
        text = normalize_text(value)
        if text and text not in seen:
            seen.add(text)
            out.append(text)
    return sep.join(out)


def make_record_id(account_id: str, seed: str = "") -> str:
    account_id = normalize_text(account_id)
    if account_id:
        safe_id = re.sub(r"[^0-9A-Za-z_-]+", "_", account_id).strip("_")
        if re.fullmatch(r"[0-9A-Za-z_-]+", account_id) and safe_id:
            return f"ACC-{safe_id}"
        digest = hashlib.sha1(account_id.encode("utf-8")).hexdigest()[:10]
        return f"ACC-{safe_id or 'TEXT'}-{digest}"
    digest = hashlib.sha1(seed.encode("utf-8")).hexdigest()[:10]
    return f"REC-{digest}"


def record_sort_key(record: dict[str, str]) -> tuple[int, int, str]:
    account_id = normalize_text(record.get("账号"))
    if record.get("记录类型") != "账户":
        return (1, 10**18, account_id)
    if account_id.isdigit():
        return (0, int(account_id), account_id)
    return (0, 10**18, account_id)


def derive_action(note: str) -> str:
    text = note.lower()
    if "限制出金" in note:
        return "限制出金"
    if "直接转T" in note or "转T" in note:
        return "P->A/T" if "放P" in note or "抛A" in note else "T"
    if "违规放T" in note or "放T" in note:
        return "T"
    if "b-p" in text or "B-P" in note:
        return "B-P"
    if "b-m" in text or "B-M" in note:
        return "B-M"
    if "m-p" in text or "M-P" in note or "放M-P" in note:
        return "M-P"
    if "放p观察" in text or "放P观察" in note:
        return "P观察"
    if "放p" in text or "放P" in note:
        return "P"
    if "放m观察" in text or "放M观察" in note:
        return "M观察"
    if "放m" in text or "放M" in note:
        return "M"
    return "待定"


def derive_group(action: str) -> str:
    if action in {"M", "M观察", "B-M"}:
        return "M"
    if action in {"P", "P观察", "B-P", "M-P", "P->A/T"}:
        return "P"
    if action == "T":
        return "T"
    if action == "限制出金":
        return "限制出金"
    return ""


def derive_tags(note: str, record_type: str) -> str:
    tag_rules = [
        ("高频", "高频"),
        ("短平", "短平"),
        ("短线", "短线"),
        ("EA", "EA"),
        ("长持仓", "长持仓"),
        ("长时间持仓", "长持仓"),
        ("大手数", "大手数"),
        ("小手数", "小手数"),
        ("逆势加仓", "逆势加仓"),
        ("稳定盈利", "稳定盈利"),
        ("盈利能力强", "盈利能力强"),
        ("抗损", "抗损"),
        ("爆仓", "爆仓风险"),
        ("同名账户", "同名账户"),
        ("跟单", "跟单/同步"),
        ("同步", "跟单/同步"),
        ("套赠金", "疑似套赠金"),
        ("credit", "Credit"),
        ("Credit", "Credit"),
        ("满仓", "满仓/过周末"),
        ("过周末", "满仓/过周末"),
        ("违规", "违规"),
        ("手数变大", "手数变大"),
        ("注册", "新注册"),
        ("没活跃", "不活跃"),
        ("IB", "IB关联"),
        ("ib", "IB关联"),
        ("穿仓", "穿仓"),
        ("菲律宾ip", "菲律宾IP"),
    ]
    tags = [label for keyword, label in tag_rules if keyword in note]
    if record_type == "IB/组":
        tags.append("IB关联")
    return uniq_join(tags)


def parse_source_records() -> list[dict[str, str]]:
    text = read_source_text()
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    records_by_id: dict[str, dict[str, str]] = {}
    last_context = ""

    for index, line in enumerate(lines, start=1):
        record_type = "账户"
        subject = ""
        m = re.match(r"^\s*((?:\d{5,10}\s*)+)(.*)$", line)
        if m:
            account_ids = re.findall(r"\d{5,10}", m.group(1))
            note = m.group(2).strip() or line
        else:
            account_ids = [""]
            note = line
            record_type = "IB/组" if line.lower().startswith("ib ") else "其他"
            ib_match = re.match(r"^(IB\s+.+?)\s+下的", line, flags=re.IGNORECASE)
            subject = ib_match.group(1) if ib_match else ""

        if note == "同上" and last_context:
            note = f"同上：{last_context}"
        elif note:
            last_context = note

        for account_id in account_ids:
            related_ids = [
                value for value in re.findall(r"\d{5,10}", note) if value != account_id
            ]
            related = uniq_join(related_ids)
            if subject:
                related = uniq_join([subject, related])
            action = derive_action(note)
            group = derive_group(action)
            tags = derive_tags(note, record_type)
            record_id = make_record_id(account_id, f"{index}:{line}")

            if record_id not in records_by_id:
                records_by_id[record_id] = {
                    "记录ID": record_id,
                    "账号": account_id,
                    "记录类型": record_type,
                    "关联账号/主体": related,
                    "建议动作": action,
                    "当前分组": group,
                    "风险标签": tags,
                    "风险/问题备注": note,
                    "原始记录": line,
                    "加入时间": datetime.now().strftime("%Y-%m-%d 12:00:00"),
                    "修改时间": now_text(),
                    "状态": "待复核",
                    "处理人/来源": SOURCE_TXT.name,
                }
            else:
                rec = records_by_id[record_id]
                rec["关联账号/主体"] = uniq_join([rec["关联账号/主体"], related])
                rec["风险标签"] = uniq_join([rec["风险标签"], tags])
                rec["风险/问题备注"] = uniq_join([rec["风险/问题备注"], note], "\n")
                rec["原始记录"] = uniq_join([rec["原始记录"], line], "\n")
                if rec["建议动作"] in {"", "待定"} and action != "待定":
                    rec["建议动作"] = action
                    rec["当前分组"] = group

    return sorted(
        records_by_id.values(),
        key=record_sort_key,
    )


def ensure_output_dir() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)


def style_accounts_sheet(ws) -> None:
    ws.freeze_panes = "B2"
    ws.sheet_view.showGridLines = False
    last_col = ws.cell(1, len(HEADERS)).column_letter
    ws.auto_filter.ref = f"A1:{last_col}{max(ws.max_row, 2)}"

    header_fill = PatternFill("solid", fgColor="263238")
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(style="thin", color="D7DEE2")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = Border(bottom=Side(style="medium", color="90A4AE"))

    widths = {
        "A": 16,
        "B": 14,
        "C": 12,
        "D": 24,
        "E": 14,
        "F": 12,
        "G": 30,
        "H": 48,
        "I": 52,
        "J": 20,
        "K": 20,
        "L": 12,
        "M": 20,
    }
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=len(HEADERS)):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=cell.column in {7, 8, 9})
            cell.border = Border(bottom=thin)
        row[0].font = Font(color="607D8B")
        row[1].number_format = "@"

    for row_idx in range(2, ws.max_row + 1):
        ws.row_dimensions[row_idx].height = 42

    add_list_validation(ws, "C", TYPE_CHOICES)
    add_list_validation(ws, "E", ACTION_CHOICES)
    add_list_validation(ws, "L", STATUS_CHOICES)

    ref = f"A1:{last_col}{max(ws.max_row, 2)}"
    if "ProblemAccounts" not in ws.tables:
        table = Table(displayName="ProblemAccounts", ref=ref)
        table.tableStyleInfo = TableStyleInfo(
            name="TableStyleMedium2",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False,
        )
        ws.add_table(table)
    else:
        ws.tables["ProblemAccounts"].ref = ref


def add_list_validation(ws, column: str, values: list[str]) -> None:
    formula = '"' + ",".join(values) + '"'
    dv = DataValidation(type="list", formula1=formula, allow_blank=True)
    ws.add_data_validation(dv)
    dv.add(f"{column}2:{column}5000")


def style_history_sheet(ws) -> None:
    ws.freeze_panes = "A2"
    ws.sheet_view.showGridLines = False
    header_fill = PatternFill("solid", fgColor="37474F")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    widths = {"A": 18, "B": 16, "C": 14, "D": 12, "E": 20, "F": 36, "G": 70, "H": 70, "I": 20}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=cell.column in {6, 7, 8})


def write_workbook(
    records: list[dict[str, str]],
    raw_lines: list[str] | None = None,
    history_rows: list[dict[str, str]] | None = None,
) -> None:
    ensure_output_dir()
    wb = Workbook()
    ws = wb.active
    ws.title = SHEET_ACCOUNTS
    ws.append(HEADERS)
    for record in records:
        ws.append([record.get(header, "") for header in HEADERS])
    style_accounts_sheet(ws)

    raw = wb.create_sheet(SHEET_RAW)
    raw.append(["序号", "原始记录"])
    for idx, line in enumerate(raw_lines or [], start=1):
        raw.append([idx, line])
    raw.column_dimensions["A"].width = 10
    raw.column_dimensions["B"].width = 100
    raw.freeze_panes = "A2"
    raw.sheet_view.showGridLines = False
    for cell in raw[1]:
        cell.fill = PatternFill("solid", fgColor="455A64")
        cell.font = Font(color="FFFFFF", bold=True)
    for row in raw.iter_rows(min_row=2, max_row=raw.max_row):
        row[1].alignment = Alignment(wrap_text=True, vertical="top")

    help_ws = wb.create_sheet(SHEET_HELP)
    help_rows = [
        ("字段", "说明"),
        ("记录ID", "网页增删改查使用的唯一键。账号记录使用 ACC-账号。"),
        ("账号", "MT4/MT5 账号，按文本原始记录拆分后去重。"),
        ("建议动作", "从原始备注初步提取，可在网页或 Excel 中调整。"),
        ("风险标签", "从备注自动抽取的可筛选标签，后续可手动增删。"),
        ("原始记录", "保留来源文本，便于回溯。"),
    ]
    for row in help_rows:
        help_ws.append(row)
    help_ws.column_dimensions["A"].width = 18
    help_ws.column_dimensions["B"].width = 80
    for cell in help_ws[1]:
        cell.fill = PatternFill("solid", fgColor="00695C")
        cell.font = Font(color="FFFFFF", bold=True)
    help_ws["B2"].comment = Comment(
        "本文件只用于本地维护已查验的问题账户记录，不执行任何 MT4/MT5 Manager 修改操作。",
        "Codex",
    )

    history = wb.create_sheet(SHEET_HISTORY)
    history.append(HISTORY_HEADERS)
    for row in history_rows or []:
        history.append([row.get(header, "") for header in HISTORY_HEADERS])
    style_history_sheet(history)

    wb.save(WORKBOOK_PATH)


def init_workbook(force: bool = False) -> None:
    if WORKBOOK_PATH.exists() and not force:
        migrate_workbook()
        return
    text = read_source_text()
    raw_lines = [line.strip() for line in text.splitlines() if line.strip()]
    write_workbook(parse_source_records(), raw_lines)
    migrate_workbook()


def read_history_rows(wb=None) -> list[dict[str, str]]:
    close_after = False
    if wb is None:
        if not WORKBOOK_PATH.exists():
            return []
        wb = load_workbook(WORKBOOK_PATH, read_only=True, data_only=True)
        close_after = True
    if SHEET_HISTORY not in wb.sheetnames:
        return []
    ws = wb[SHEET_HISTORY]
    headers = [normalize_text(cell.value) for cell in ws[1]]
    rows: list[dict[str, str]] = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        item = {headers[i]: normalize_text(row[i]) for i in range(min(len(headers), len(row)))}
        for header in HISTORY_HEADERS:
            item.setdefault(header, "")
        rows.append(item)
    if close_after:
        wb.close()
    return rows


def history_changed_fields(before: dict[str, str], after: dict[str, str]) -> str:
    changed = []
    for header in HEADERS:
        if header == "修改时间":
            continue
        if normalize_text(before.get(header)) != normalize_text(after.get(header)):
            changed.append(header)
    return "；".join(changed)


def make_history_row(
    before: dict[str, str],
    after: dict[str, str] | None,
    operation: str,
    changed_fields: str | None = None,
) -> dict[str, str]:
    after = after or {}
    stamp = now_text()
    return {
        "历史ID": f"HIS-{datetime.now().strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:8]}",
        "记录ID": before.get("记录ID") or after.get("记录ID", ""),
        "账号": before.get("账号") or after.get("账号", ""),
        "操作": operation,
        "修改时间": stamp,
        "修改字段": changed_fields if changed_fields is not None else history_changed_fields(before, after),
        "修改前JSON": json.dumps(before, ensure_ascii=False, sort_keys=True),
        "修改后JSON": json.dumps(after, ensure_ascii=False, sort_keys=True) if after else "",
        "处理人/来源": after.get("处理人/来源") or before.get("处理人/来源", ""),
    }


def migrate_workbook() -> None:
    if not WORKBOOK_PATH.exists():
        return
    wb = load_workbook(WORKBOOK_PATH)
    if SHEET_ACCOUNTS not in wb.sheetnames:
        return
    ws = wb[SHEET_ACCOUNTS]
    current_headers = [normalize_text(cell.value) for cell in ws[1]]
    needs_migration = current_headers != HEADERS or SHEET_HISTORY not in wb.sheetnames
    if not needs_migration:
        return

    header_index = {header: idx for idx, header in enumerate(current_headers)}
    journal_times = account_join_times_from_journals()
    records: list[dict[str, str]] = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        rec: dict[str, str] = {}
        for header in HEADERS:
            value = ""
            for source_header in LEGACY_HEADER_ALIASES.get(header, [header]):
                if source_header in header_index and header_index[source_header] < len(row):
                    value = normalize_text(row[header_index[source_header]])
                    if value:
                        break
            rec[header] = value
        if not rec["记录ID"]:
            rec["记录ID"] = make_record_id(rec["账号"], json.dumps(rec, ensure_ascii=False))
        rec["加入时间"] = default_join_time(rec, journal_times)
        rec["修改时间"] = normalize_datetime(rec["修改时间"]) or rec["加入时间"]
        records.append(rec)

    raw_lines: list[str] = []
    if SHEET_RAW in wb.sheetnames:
        raw_ws = wb[SHEET_RAW]
        for row in raw_ws.iter_rows(min_row=2, values_only=True):
            if len(row) > 1 and row[1]:
                raw_lines.append(str(row[1]))
    history_rows = read_history_rows(wb)
    write_workbook(records, raw_lines, history_rows)


def normalize_initial_record_times() -> None:
    if not WORKBOOK_PATH.exists():
        return
    records = load_records()
    history_rows = read_history_rows()
    changed_ids = {row["记录ID"] for row in history_rows if row["操作"] == "修改"}
    journal_times = account_join_times_from_journals()
    changed = False
    for record in records:
        joined_at = default_join_time(record, journal_times)
        if record["加入时间"] != joined_at:
            record["加入时间"] = joined_at
            changed = True
        if record["记录ID"] not in changed_ids and record["修改时间"] != record["加入时间"]:
            record["修改时间"] = record["加入时间"]
            changed = True
    if changed:
        save_records(records)


def load_records() -> list[dict[str, str]]:
    init_workbook()
    wb = load_workbook(WORKBOOK_PATH)
    ws = wb[SHEET_ACCOUNTS]
    headers = [normalize_text(cell.value) for cell in ws[1]]
    header_index = {header: idx for idx, header in enumerate(headers)}
    records: list[dict[str, str]] = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        rec = {}
        for header in HEADERS:
            value = ""
            for source_header in LEGACY_HEADER_ALIASES.get(header, [header]):
                idx = header_index.get(source_header)
                if idx is not None and idx < len(row):
                    value = normalize_text(row[idx])
                    if value:
                        break
            rec[header] = value
        if not rec["记录ID"]:
            rec["记录ID"] = make_record_id(rec["账号"], json.dumps(rec, ensure_ascii=False))
        rec["加入时间"] = normalize_datetime(rec["加入时间"])
        rec["修改时间"] = normalize_datetime(rec["修改时间"]) or rec["加入时间"]
        records.append(rec)
    wb.close()
    return records


def save_records(records: list[dict[str, str]], extra_history: list[dict[str, str]] | None = None) -> None:
    raw_lines: list[str] = []
    history_rows: list[dict[str, str]] = []
    if WORKBOOK_PATH.exists():
        wb = load_workbook(WORKBOOK_PATH, read_only=True, data_only=True)
        if SHEET_RAW in wb.sheetnames:
            raw_ws = wb[SHEET_RAW]
            for row in raw_ws.iter_rows(min_row=2, values_only=True):
                if len(row) > 1 and row[1]:
                    raw_lines.append(str(row[1]))
        history_rows = read_history_rows(wb)
        wb.close()
    history_rows.extend(extra_history or [])
    write_workbook(records, raw_lines, history_rows)


def summarize(records: list[dict[str, str]]) -> dict:
    actions: dict[str, int] = {}
    statuses: dict[str, int] = {}
    for record in records:
        actions[record["建议动作"] or "未填"] = actions.get(record["建议动作"] or "未填", 0) + 1
        statuses[record["状态"] or "未填"] = statuses.get(record["状态"] or "未填", 0) + 1
    return {
        "total": len(records),
        "accountRecords": sum(1 for r in records if r["账号"]),
        "groupRecords": sum(1 for r in records if not r["账号"]),
        "actions": actions,
        "statuses": statuses,
        "workbook": str(WORKBOOK_PATH),
        "updatedAt": now_text(),
    }


def public_chart_url(path: Path) -> str:
    return "/chart-file/" + quote(path.name)


def parse_chart_file(path: Path) -> dict[str, str | int | bool]:
    name = path.name
    stem = name[: -len("_trade_kline.html")] if name.endswith("_trade_kline.html") else path.stem
    parts = stem.split("_")
    account = parts[0] if parts else ""
    start = ""
    end = ""
    if len(parts) >= 5:
        start = f"{parts[1]}_{parts[2]}"
        end = f"{parts[3]}_{parts[4]}"
    stat = path.stat()
    return {
        "account": account,
        "name": name,
        "path": str(path),
        "url": public_chart_url(path),
        "size": stat.st_size,
        "sizeText": f"{stat.st_size / 1024 / 1024:.1f} MB",
        "mtime": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
        "start": start,
        "end": end,
    }


def scan_chart_files(records: list[dict[str, str]] | None = None) -> list[dict]:
    records = records or load_records()
    known_accounts = {record["账号"] for record in records if record["账号"]}
    charts = []
    for path in sorted(KLINE_OUT_DIR.glob("*_trade_kline.html"), key=lambda p: p.stat().st_mtime, reverse=True):
        item = parse_chart_file(path)
        item["inRegistry"] = item["account"] in known_accounts
        item["recordId"] = ""
        for record in records:
            if record["账号"] == item["account"]:
                item["recordId"] = record["记录ID"]
                item["status"] = record["状态"]
                item["action"] = record["建议动作"]
                break
        charts.append(item)
    return charts


def is_today_text(value: str, today: str | None = None) -> bool:
    today = today or datetime.now().strftime("%Y-%m-%d")
    return normalize_text(value).startswith(today)


def daily_report_records() -> list[dict[str, str]]:
    today = datetime.now().strftime("%Y-%m-%d")
    records = load_records()
    history_rows = read_history_rows()
    today_ids = {
        record["记录ID"]
        for record in records
        if is_today_text(record["加入时间"], today) or is_today_text(record["修改时间"], today)
    }
    today_ids.update(
        row["记录ID"]
        for row in history_rows
        if is_today_text(row["修改时间"], today) and row["操作"] != "删除"
    )
    selected = [record for record in records if record["记录ID"] in today_ids]
    return sorted(selected, key=lambda record: (record["加入时间"], record["账号"] or record["关联账号/主体"]))


def daily_report_docx_bytes() -> tuple[str, bytes]:
    today = datetime.now()
    records = daily_report_records()
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"问题账户日报 {today.strftime('%m%d')}")
    run.bold = True
    run.font.size = Pt(16)

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["账号", "建议", "备注"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for record in records:
        row = table.add_row().cells
        row[0].text = record["账号"] or record["关联账号/主体"] or record["记录ID"]
        row[1].text = record["建议动作"] or ""
        row[2].text = record["风险/问题备注"] or ""

    output_path = OUT_DIR / f"journal_{today.strftime('%m%d')}.docx"
    document.save(output_path)
    return output_path.name, output_path.read_bytes()


def action_choices_for(records: list[dict[str, str]]) -> list[str]:
    choices = list(ACTION_CHOICES)
    for record in records:
        action = normalize_text(record.get("建议动作"))
        if action and action not in choices:
            insert_at = max(len(choices) - 2, 0)
            choices.insert(insert_at, action)
    return choices


def json_response(handler: BaseHTTPRequestHandler, payload: object, status: int = 200) -> None:
    data = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Cache-Control", "no-store")
    handler.send_header("Content-Length", str(len(data)))
    handler.end_headers()
    handler.wfile.write(data)


def html_response(handler: BaseHTTPRequestHandler, body: str, status: int = 200) -> None:
    data = body.encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "text/html; charset=utf-8")
    handler.send_header("Cache-Control", "no-store")
    handler.send_header("Content-Length", str(len(data)))
    handler.end_headers()
    handler.wfile.write(data)


def error_response(handler: BaseHTTPRequestHandler, message: str, status: int = 400) -> None:
    json_response(handler, {"ok": False, "error": message}, status)


def parse_body(handler: BaseHTTPRequestHandler) -> dict:
    length = int(handler.headers.get("Content-Length", "0") or "0")
    raw = handler.rfile.read(length) if length else b"{}"
    return json.loads(raw.decode("utf-8") or "{}")


def normalize_payload(payload: dict, existing: dict[str, str] | None = None) -> dict[str, str]:
    existing = existing or {}
    rec = {header: normalize_text(existing.get(header, "")) for header in HEADERS}
    for field in EDITABLE_FIELDS:
        if field in payload:
            rec[field] = normalize_text(payload[field])
    if not rec["记录类型"]:
        rec["记录类型"] = "账户" if rec["账号"] else "其他"
    if not rec["建议动作"]:
        rec["建议动作"] = derive_action(rec["风险/问题备注"])
    if not rec["当前分组"]:
        rec["当前分组"] = derive_group(rec["建议动作"])
    if not rec["风险标签"]:
        rec["风险标签"] = derive_tags(rec["风险/问题备注"], rec["记录类型"])
    if not rec["状态"]:
        rec["状态"] = "待复核"
    if not rec["加入时间"]:
        rec["加入时间"] = now_text()
    else:
        rec["加入时间"] = normalize_datetime(rec["加入时间"])
    rec["修改时间"] = now_text()
    rec["记录ID"] = existing.get("记录ID") or make_record_id(
        rec["账号"], f"{rec['关联账号/主体']}:{rec['风险/问题备注']}:{uuid.uuid4().hex}"
    )
    return rec


class Handler(BaseHTTPRequestHandler):
    server_version = "AccountRegistry/1.0"

    def log_message(self, fmt: str, *args) -> None:
        sys.stderr.write("[%s] %s\n" % (now_text(), fmt % args))

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        path = parsed.path
        if path == "/":
            html_response(self, INDEX_HTML)
            return
        if path == "/api/accounts":
            records = load_records()
            history_counts: dict[str, int] = {}
            for row in read_history_rows():
                if row["操作"] == "修改":
                    history_counts[row["记录ID"]] = history_counts.get(row["记录ID"], 0) + 1
            json_response(
                self,
                {
                    "ok": True,
                    "summary": summarize(records),
                    "records": records,
                    "historyCounts": history_counts,
                },
            )
            return
        if path == "/api/charts":
            records = load_records()
            charts = scan_chart_files(records)
            json_response(
                self,
                {
                    "ok": True,
                    "charts": charts,
                    "summary": {
                        "total": len(charts),
                        "linked": sum(1 for chart in charts if chart.get("inRegistry")),
                        "unlinked": sum(1 for chart in charts if not chart.get("inRegistry")),
                        "uploadUrl": TRADE_KLINE_WEB_URL,
                    },
                },
            )
            return
        chart_match = re.match(r"^/api/accounts/(.+)/charts$", path)
        if chart_match:
            record_id = unquote(chart_match.group(1))
            records = load_records()
            record = next((row for row in records if row["记录ID"] == record_id), None)
            if not record:
                error_response(self, "记录不存在", 404)
                return
            charts = [
                chart
                for chart in scan_chart_files(records)
                if chart["account"] == record["账号"]
            ]
            json_response(self, {"ok": True, "record": record, "charts": charts})
            return
        history_match = re.match(r"^/api/accounts/(.+)/history$", path)
        if history_match:
            record_id = unquote(history_match.group(1))
            rows = [row for row in read_history_rows() if row["记录ID"] == record_id]
            rows.sort(key=lambda row: row["修改时间"], reverse=True)
            json_response(self, {"ok": True, "history": rows})
            return
        if path == "/api/meta":
            records = load_records()
            json_response(
                self,
                {
                    "ok": True,
                    "summary": summarize(records),
                    "actions": action_choices_for(records),
                    "statuses": STATUS_CHOICES,
                    "types": TYPE_CHOICES,
                },
            )
            return
        if path == "/download/problematic_accounts.xlsx":
            init_workbook()
            data = WORKBOOK_PATH.read_bytes()
            self.send_response(200)
            self.send_header(
                "Content-Type",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            self.send_header(
                "Content-Disposition",
                f"attachment; filename*=UTF-8''{quote(WORKBOOK_PATH.name)}",
            )
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        if path == "/download/daily-report":
            filename, data = daily_report_docx_bytes()
            self.send_response(200)
            self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            self.send_header(
                "Content-Disposition",
                f"attachment; filename*=UTF-8''{quote(filename)}",
            )
            self.send_header("Cache-Control", "no-store")
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        file_match = re.match(r"^/chart-file/(.+)$", path)
        if file_match:
            name = Path(unquote(file_match.group(1))).name
            file_path = (KLINE_OUT_DIR / name).resolve()
            if file_path.parent != KLINE_OUT_DIR.resolve() or not file_path.exists() or not name.endswith("_trade_kline.html"):
                error_response(self, "图表文件不存在", 404)
                return
            data = file_path.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Cache-Control", "no-store")
            self.send_header("Content-Length", str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        error_response(self, "Not found", 404)

    def do_POST(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/api/accounts":
            payload = parse_body(self)
            records = load_records()
            rec = normalize_payload(payload)
            if rec["账号"] and any(r["账号"] == rec["账号"] for r in records):
                error_response(self, f"账号 {rec['账号']} 已存在")
                return
            records.append(rec)
            records.sort(key=record_sort_key)
            save_records(records)
            json_response(self, {"ok": True, "record": rec, "summary": summarize(records)})
            return
        if parsed.path == "/api/rebuild":
            init_workbook(force=True)
            records = load_records()
            json_response(self, {"ok": True, "summary": summarize(records), "records": records})
            return
        error_response(self, "Not found", 404)

    def do_PUT(self) -> None:
        parsed = urlparse(self.path)
        match = re.match(r"^/api/accounts/(.+)$", parsed.path)
        if not match:
            error_response(self, "Not found", 404)
            return
        record_id = unquote(match.group(1))
        payload = parse_body(self)
        records = load_records()
        for idx, record in enumerate(records):
            if record["记录ID"] == record_id:
                updated = normalize_payload(payload, record)
                if updated["账号"] and any(
                    r["账号"] == updated["账号"] and r["记录ID"] != record_id for r in records
                ):
                    error_response(self, f"账号 {updated['账号']} 已存在")
                    return
                records[idx] = updated
                history = []
                changed = history_changed_fields(record, updated)
                if changed:
                    history.append(make_history_row(record, updated, "修改", changed))
                save_records(records, history)
                json_response(self, {"ok": True, "record": updated, "summary": summarize(records)})
                return
        error_response(self, "记录不存在", 404)

    def do_DELETE(self) -> None:
        parsed = urlparse(self.path)
        match = re.match(r"^/api/accounts/(.+)$", parsed.path)
        if not match:
            error_response(self, "Not found", 404)
            return
        record_id = unquote(match.group(1))
        records = load_records()
        kept = [record for record in records if record["记录ID"] != record_id]
        if len(kept) == len(records):
            error_response(self, "记录不存在", 404)
            return
        removed = [record for record in records if record["记录ID"] == record_id]
        history = [make_history_row(removed[0], None, "删除", "删除记录")] if removed else []
        save_records(kept, history)
        json_response(self, {"ok": True, "summary": summarize(kept)})


INDEX_HTML = r"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>问题账户台账</title>
  <style>
    :root {
      --bg: #f6f8f9;
      --panel: #ffffff;
      --ink: #1f2933;
      --muted: #667783;
      --line: #d9e1e5;
      --teal: #00796b;
      --teal-soft: #e0f2ef;
      --red: #c62828;
      --amber: #f2a900;
      --green: #2e7d32;
      --gray: #455a64;
      --focus: #0097a7;
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      background: var(--bg);
      color: var(--ink);
      font-family: "Microsoft YaHei", "Segoe UI", Arial, sans-serif;
    }
    header {
      height: 64px;
      display: flex;
      align-items: center;
      justify-content: space-between;
      padding: 0 24px;
      background: #ffffff;
      border-bottom: 1px solid var(--line);
    }
    h1 { margin: 0; font-size: 22px; font-weight: 700; letter-spacing: 0; }
    main { padding: 18px 24px 24px; }
    .summary {
      display: grid;
      grid-template-columns: repeat(4, minmax(160px, 1fr));
      gap: 12px;
      margin-bottom: 14px;
    }
    .metric {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 12px 14px;
      min-height: 74px;
    }
    .metric b { display: block; font-size: 24px; margin-top: 4px; }
    .metric span { color: var(--muted); font-size: 13px; }
    .toolbar {
      display: grid;
      grid-template-columns: minmax(220px, 1fr) 150px 150px 170px auto auto auto auto;
      gap: 10px;
      align-items: center;
      margin-bottom: 12px;
    }
    input, select, textarea {
      width: 100%;
      border: 1px solid var(--line);
      border-radius: 6px;
      background: #fff;
      color: var(--ink);
      font: inherit;
      padding: 9px 10px;
      outline: none;
    }
    input:focus, select:focus, textarea:focus {
      border-color: var(--focus);
      box-shadow: 0 0 0 3px rgba(0, 151, 167, .14);
    }
    button, .button-link {
      border: 1px solid transparent;
      border-radius: 6px;
      padding: 9px 12px;
      font: inherit;
      cursor: pointer;
      background: var(--gray);
      color: #fff;
      text-decoration: none;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      gap: 6px;
      min-height: 38px;
      white-space: nowrap;
    }
    button.primary { background: var(--teal); }
    button.danger { background: var(--red); }
    button.light, .button-link.light {
      background: #fff;
      color: var(--ink);
      border-color: var(--line);
    }
    .layout {
      display: grid;
      grid-template-columns: minmax(0, 1fr) 400px;
      gap: 14px;
      align-items: start;
    }
    .table-wrap, .editor, .chart-panel {
      background: var(--panel);
      border: 1px solid var(--line);
      border-radius: 8px;
      overflow: hidden;
    }
    table {
      width: 100%;
      border-collapse: collapse;
      table-layout: fixed;
    }
    th, td {
      border-bottom: 1px solid var(--line);
      padding: 10px;
      vertical-align: top;
      font-size: 14px;
    }
    th {
      position: sticky;
      top: 0;
      z-index: 1;
      background: #263238;
      color: #fff;
      text-align: left;
      font-weight: 700;
    }
    tr:hover td { background: #f4fbfa; }
    .col-account { width: 128px; }
    .col-action { width: 96px; }
    .col-tags { width: 260px; }
    .col-time { width: 170px; }
    .col-status { width: 88px; }
    .col-tools { width: 128px; text-align: right; }
    .row-actions { display: flex; gap: 6px; justify-content: flex-end; flex-wrap: wrap; }
    .row-actions button { min-height: 34px; padding: 6px 10px; }
    .note {
      display: -webkit-box;
      -webkit-line-clamp: 3;
      -webkit-box-orient: vertical;
      overflow: hidden;
      line-height: 1.5;
    }
    .tags {
      display: flex;
      flex-wrap: wrap;
      gap: 5px;
    }
    .tag {
      display: inline-flex;
      align-items: center;
      max-width: 100%;
      padding: 2px 7px;
      border-radius: 999px;
      background: var(--teal-soft);
      color: #00695c;
      font-size: 12px;
      line-height: 18px;
    }
    .pill {
      display: inline-flex;
      align-items: center;
      min-height: 24px;
      padding: 2px 8px;
      border-radius: 999px;
      background: #eceff1;
      color: #37474f;
      font-size: 12px;
      font-weight: 700;
    }
    .pill.p { background: #ffebee; color: #b71c1c; }
    .pill.m { background: #e8f5e9; color: #1b5e20; }
    .pill.t { background: #fff8e1; color: #8a5a00; }
    .editor { padding: 14px; position: sticky; top: 82px; }
    .editor h2 { margin: 0 0 12px; font-size: 18px; }
    .form-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 10px; }
    .form-grid label { display: block; color: var(--muted); font-size: 12px; margin-bottom: 5px; }
    .span-2 { grid-column: 1 / -1; }
    textarea { min-height: 86px; resize: vertical; line-height: 1.45; }
    .editor-actions { display: flex; gap: 8px; justify-content: flex-end; margin-top: 12px; }
    .statusbar { min-height: 22px; margin-top: 10px; color: var(--muted); font-size: 13px; }
    .empty { padding: 36px; color: var(--muted); text-align: center; }
    .time-text { color: var(--muted); font-size: 12px; line-height: 1.45; }
    .list-time { font-size: 13px; line-height: 1.5; white-space: normal; word-break: keep-all; }
    .list-time b { color: var(--ink); font-weight: 700; }
    dialog {
      width: min(880px, calc(100vw - 32px));
      max-height: 82vh;
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 0;
      overflow: hidden;
    }
    dialog::backdrop { background: rgba(31, 41, 51, .32); }
    .history-head {
      display: flex;
      justify-content: space-between;
      align-items: center;
      gap: 10px;
      padding: 14px 16px;
      border-bottom: 1px solid var(--line);
    }
    .history-body {
      max-height: 68vh;
      overflow: auto;
      padding: 12px 16px 16px;
      background: #f8fafb;
    }
    .history-item {
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 0;
      margin-bottom: 12px;
      background: #fff;
      overflow: hidden;
    }
    .version-head {
      display: flex;
      justify-content: space-between;
      gap: 10px;
      align-items: flex-start;
      padding: 12px 14px;
      border-bottom: 1px solid var(--line);
      background: #fbfdfe;
    }
    .version-title { display: flex; flex-wrap: wrap; gap: 8px; align-items: center; }
    .version-badge {
      display: inline-flex;
      align-items: center;
      min-height: 24px;
      padding: 2px 9px;
      border-radius: 999px;
      background: var(--teal-soft);
      color: #00695c;
      font-weight: 700;
      font-size: 12px;
    }
    .version-badge.delete { background: #ffebee; color: #b71c1c; }
    .version-time { color: var(--muted); font-size: 12px; line-height: 1.5; text-align: right; }
    .version-body { display: grid; grid-template-columns: 1fr; gap: 12px; padding: 12px 14px 14px; }
    .version-panel {
      border: 1px solid var(--line);
      border-radius: 8px;
      overflow: hidden;
      background: #fff;
    }
    .version-panel h4 {
      margin: 0;
      padding: 8px 10px;
      background: #eceff1;
      font-size: 13px;
    }
    .version-panel.after h4 { background: #e0f2ef; color: #00695c; }
    .version-panel.before h4 { background: #fff8e1; color: #7a5600; }
    .version-grid { display: grid; grid-template-columns: 92px minmax(0, 1fr); }
    .version-label, .version-value {
      padding: 8px 10px;
      border-top: 1px solid #edf1f3;
      font-size: 13px;
      line-height: 1.45;
      min-width: 0;
    }
    .version-label { color: var(--muted); background: #fafafa; }
    .version-value { white-space: pre-wrap; word-break: break-word; }
    .version-value.changed {
      background: #fffde7;
      box-shadow: inset 3px 0 0 var(--amber);
    }
    .pagination {
      display: flex;
      justify-content: space-between;
      align-items: center;
      gap: 12px;
      padding: 12px 14px;
      background: #fff;
      border-top: 1px solid var(--line);
      flex-wrap: wrap;
    }
    .pager-buttons { display: flex; gap: 8px; align-items: center; }
    .pager-buttons button { min-height: 34px; padding: 6px 10px; }
    .pager-info { color: var(--muted); font-size: 13px; }
    @media (max-width: 820px) {
      .version-time { text-align: left; }
    }
    .chart-section { margin-top: 14px; border-top: 1px solid var(--line); padding-top: 12px; }
    .section-title { display: flex; justify-content: space-between; align-items: center; gap: 8px; margin-bottom: 8px; }
    .section-title b { font-size: 15px; }
    .chart-list { display: grid; gap: 8px; max-height: 220px; overflow: auto; padding-right: 2px; }
    .chart-item {
      border: 1px solid var(--line);
      border-radius: 8px;
      padding: 9px;
      background: #fff;
      display: grid;
      gap: 7px;
    }
    .chart-item.active { border-color: var(--teal); box-shadow: 0 0 0 3px rgba(0, 121, 107, .12); }
    .chart-name { font-weight: 700; word-break: break-all; font-size: 13px; }
    .chart-meta { color: var(--muted); font-size: 12px; line-height: 1.45; }
    .chart-actions { display: flex; flex-wrap: wrap; gap: 6px; }
    .chart-actions a, .chart-actions button { min-height: 32px; padding: 6px 9px; font-size: 13px; }
    .chart-frame {
      width: 100%;
      height: 360px;
      border: 1px solid var(--line);
      border-radius: 8px;
      margin-top: 10px;
      background: #fff;
    }
    .muted { color: var(--muted); }
    @media (max-width: 1100px) {
      .summary { grid-template-columns: repeat(2, minmax(140px, 1fr)); }
      .layout { grid-template-columns: 1fr; }
      .editor { position: static; }
      .toolbar { grid-template-columns: 1fr 1fr; }
    }
    @media (max-width: 680px) {
      header { height: auto; gap: 10px; padding: 14px; flex-wrap: wrap; }
      main { padding: 14px; }
      .summary { grid-template-columns: 1fr; }
      .toolbar { grid-template-columns: 1fr; }
      .form-grid { grid-template-columns: 1fr; }
      .span-2 { grid-column: auto; }
      th:nth-child(4), td:nth-child(4),
      th:nth-child(5), td:nth-child(5),
      th:nth-child(6), td:nth-child(6) { display: none; }
    }
  </style>
</head>
<body>
  <header>
    <h1>问题账户台账</h1>
    <a class="button-link light" href="/download/problematic_accounts.xlsx">下载 Excel</a>
  </header>
  <main>
    <section class="summary">
      <div class="metric"><span>记录数</span><b id="total">0</b></div>
      <div class="metric"><span>账号记录</span><b id="accountRecords">0</b></div>
      <div class="metric"><span>IB/组记录</span><b id="groupRecords">0</b></div>
      <div class="metric"><span>更新时间</span><b id="updatedAt" style="font-size:15px">-</b></div>
    </section>
    <section class="toolbar">
      <input id="search" placeholder="搜索账号、标签、备注、IB" />
      <select id="actionFilter"></select>
      <select id="statusFilter"></select>
      <select id="sortBy">
        <option value="joined_desc">加入时间 新到旧</option>
        <option value="account">按账号排序</option>
        <option value="joined_asc">加入时间 旧到新</option>
        <option value="updated_desc">更新时间 新到旧</option>
        <option value="updated_asc">更新时间 旧到新</option>
      </select>
      <button class="primary" id="newBtn">＋ 新增</button>
      <button class="light" id="chartLibraryBtn">图表库</button>
      <a class="button-link light" href="http://127.0.0.1:8765" target="_blank" rel="noopener">生成图表</a>
      <a class="button-link light" href="/download/daily-report">导出日报</a>
      <button class="light" id="reloadBtn">刷新</button>
    </section>
    <section class="layout">
      <div class="table-wrap">
        <table>
          <thead>
            <tr>
              <th class="col-account">账号</th>
              <th class="col-action">建议</th>
              <th class="col-tags">标签</th>
              <th>备注</th>
              <th class="col-time">时间</th>
              <th class="col-status">状态</th>
              <th class="col-tools"></th>
            </tr>
          </thead>
          <tbody id="rows"></tbody>
        </table>
        <div class="empty" id="empty" hidden>没有匹配记录</div>
        <div class="pagination">
          <div class="pager-info" id="pageInfo">第 1 页</div>
          <div class="pager-buttons">
            <button class="light" id="prevPageBtn">上一页</button>
            <button class="light" id="nextPageBtn">下一页</button>
          </div>
        </div>
      </div>
      <aside class="editor">
        <h2 id="formTitle">新增记录</h2>
        <div class="form-grid">
          <div>
            <label>账号</label>
            <input id="account" />
          </div>
          <div>
            <label>记录类型</label>
            <select id="type"></select>
          </div>
          <div>
            <label>建议动作</label>
      <select id="action"></select>
            <input id="customAction" placeholder="输入自定义建议动作" style="margin-top:8px;display:none" />
          </div>
          <div>
            <label>当前分组</label>
            <input id="group" />
          </div>
          <div class="span-2">
            <label>关联账号/主体</label>
            <input id="related" />
          </div>
          <div class="span-2">
            <label>风险标签</label>
            <input id="tags" />
          </div>
          <div class="span-2">
            <label>风险/问题备注</label>
            <textarea id="note"></textarea>
          </div>
          <div class="span-2">
            <label>原始记录</label>
            <textarea id="raw"></textarea>
          </div>
          <div>
            <label>状态</label>
            <select id="status"></select>
          </div>
          <div>
            <label>处理人/来源</label>
            <input id="source" />
          </div>
          <div>
            <label>加入时间</label>
            <input id="joinedAt" readonly />
          </div>
          <div>
            <label>修改时间</label>
            <input id="updatedAtField" readonly />
          </div>
        </div>
        <div class="editor-actions">
          <button class="light" id="clearBtn">清空</button>
          <button class="light" id="historyBtn" hidden>历史</button>
          <button class="danger" id="deleteBtn" hidden>删除</button>
          <button class="primary" id="saveBtn">保存</button>
        </div>
        <div class="statusbar" id="statusbar"></div>
        <div class="chart-section">
          <div class="section-title">
            <b>匹配图表</b>
            <span class="muted" id="chartHint">选择账号后显示</span>
          </div>
          <div class="chart-list" id="accountCharts"></div>
          <iframe class="chart-frame" id="chartPreview" title="买卖点图预览" hidden></iframe>
        </div>
      </aside>
    </section>
    <dialog id="historyDialog">
      <div class="history-head">
        <b id="historyTitle">修改历史</b>
        <button class="light" id="closeHistoryBtn">关闭</button>
      </div>
      <div class="history-body" id="historyBody"></div>
    </dialog>
    <dialog id="chartLibraryDialog">
      <div class="history-head">
        <b id="chartLibraryTitle">图表库</b>
        <button class="light" id="closeChartLibraryBtn">关闭</button>
      </div>
      <div class="history-body">
        <div class="toolbar" style="grid-template-columns: minmax(180px, 1fr) 140px auto; padding:0; margin-bottom:10px">
          <input id="chartSearch" placeholder="搜索账号或图表文件" />
          <select id="chartLinkedFilter">
            <option value="">全部图表</option>
            <option value="linked">已入台账</option>
            <option value="unlinked">未入台账</option>
          </select>
          <a class="button-link light" href="http://127.0.0.1:8765" target="_blank" rel="noopener">上传生成</a>
        </div>
        <div class="chart-list" id="chartLibraryList" style="max-height:60vh"></div>
      </div>
    </dialog>
  </main>
  <script>
    const state = { records: [], meta: {}, editingId: null, historyCounts: {}, accountCharts: [], allCharts: [], page: 1, pageSize: 20, lastFilteredCount: 0 };
    const $ = (id) => document.getElementById(id);
    const fields = {
      "账号": "account",
      "记录类型": "type",
      "关联账号/主体": "related",
      "建议动作": "action",
      "当前分组": "group",
      "风险标签": "tags",
      "风险/问题备注": "note",
      "原始记录": "raw",
      "状态": "status",
      "处理人/来源": "source",
    };

    function setStatus(text, good = true) {
      $("statusbar").textContent = text || "";
      $("statusbar").style.color = good ? "var(--muted)" : "var(--red)";
    }
    function fillSelect(select, values, first) {
      select.innerHTML = "";
      if (first !== undefined) select.append(new Option(first, ""));
      values.forEach((value) => select.append(new Option(value || "未填", value)));
    }
    function ensureSelectOption(select, value) {
      const text = String(value || "").trim();
      if (!text) return;
      if (!Array.from(select.options).some((opt) => opt.value === text)) {
        const customIndex = Array.from(select.options).findIndex((opt) => opt.value === "自定义");
        const option = new Option(text, text);
        if (customIndex >= 0) select.add(option, select.options[customIndex]);
        else select.append(option);
      }
    }
    function actionClass(action) {
      if ((action || "").includes("T")) return "t";
      if ((action || "").includes("P")) return "p";
      if ((action || "").includes("M")) return "m";
      return "";
    }
    function escapeText(text) {
      return String(text || "").replace(/[&<>"']/g, (c) => ({
        "&": "&amp;", "<": "&lt;", ">": "&gt;", '"': "&quot;", "'": "&#39;"
      }[c]));
    }
    function splitTags(text) {
      return String(text || "").split(/[；;,，\n]+/).map((x) => x.trim()).filter(Boolean);
    }
    function visibleTime(record) {
      const hasHistory = (state.historyCounts[record["记录ID"]] || 0) > 0;
      return {
        label: hasHistory ? "修改" : "加入",
        value: hasHistory ? (record["修改时间"] || record["加入时间"] || "-") : (record["加入时间"] || "-"),
      };
    }
    function effectiveTimeValue(record) {
      return visibleTime(record).value;
    }
    function renderSummary(summary) {
      $("total").textContent = summary.total || 0;
      $("accountRecords").textContent = summary.accountRecords || 0;
      $("groupRecords").textContent = summary.groupRecords || 0;
      $("updatedAt").textContent = summary.updatedAt || "-";
    }
    function filteredRecords() {
      const q = $("search").value.trim().toLowerCase();
      const action = $("actionFilter").value;
      const status = $("statusFilter").value;
      const rows = state.records.filter((record) => {
        const hay = Object.values(record).join(" ").toLowerCase();
        return (!q || hay.includes(q)) &&
          (!action || record["建议动作"] === action) &&
          (!status || record["状态"] === status);
      });
      const asTime = (value) => {
        const text = String(value || "").trim();
        if (!text) return 0;
        const normalized = text.length <= 10 ? `${text} 00:00:00` : text;
        const parsed = Date.parse(normalized.replace(/-/g, "/"));
        return Number.isNaN(parsed) ? 0 : parsed;
      };
      const asAccount = (value) => {
        const parsed = Number.parseInt(String(value || "").replace(/\D/g, ""), 10);
        return Number.isNaN(parsed) ? Number.MAX_SAFE_INTEGER : parsed;
      };
      const sortBy = $("sortBy").value;
      rows.sort((a, b) => {
        if (sortBy === "joined_desc") return asTime(effectiveTimeValue(b)) - asTime(effectiveTimeValue(a));
        if (sortBy === "joined_asc") return asTime(a["加入时间"]) - asTime(b["加入时间"]);
        if (sortBy === "updated_desc") return asTime(b["修改时间"]) - asTime(a["修改时间"]);
        if (sortBy === "updated_asc") return asTime(a["修改时间"]) - asTime(b["修改时间"]);
        return asAccount(a["账号"]) - asAccount(b["账号"]);
      });
      return rows;
    }
    function updatePagination(total) {
      const pages = Math.max(1, Math.ceil(total / state.pageSize));
      if (state.page > pages) state.page = pages;
      const start = total ? (state.page - 1) * state.pageSize + 1 : 0;
      const end = Math.min(total, state.page * state.pageSize);
      $("pageInfo").textContent = `第 ${state.page} / ${pages} 页 · ${start}-${end} / ${total}`;
      $("prevPageBtn").disabled = state.page <= 1;
      $("nextPageBtn").disabled = state.page >= pages;
    }
    function renderRows() {
      const tbody = $("rows");
      const rows = filteredRecords();
      state.lastFilteredCount = rows.length;
      updatePagination(rows.length);
      const pageRows = rows.slice((state.page - 1) * state.pageSize, state.page * state.pageSize);
      tbody.innerHTML = pageRows.map((record) => {
        const tags = splitTags(record["风险标签"]).map((tag) => `<span class="tag">${escapeText(tag)}</span>`).join("");
        const account = record["账号"] || record["关联账号/主体"] || "未填";
        const time = visibleTime(record);
        return `<tr>
          <td class="col-account"><b>${escapeText(account)}</b><br><span style="color:var(--muted);font-size:12px">${escapeText(record["记录类型"])}</span></td>
          <td class="col-action"><span class="pill ${actionClass(record["建议动作"])}">${escapeText(record["建议动作"] || "待定")}</span></td>
          <td class="col-tags"><div class="tags">${tags || '<span style="color:var(--muted)">-</span>'}</div></td>
          <td><div class="note">${escapeText(record["风险/问题备注"])}</div></td>
          <td class="col-time"><div class="list-time"><b>${escapeText(time.label)}</b><br>${escapeText(time.value)}</div></td>
          <td class="col-status">${escapeText(record["状态"])}</td>
          <td class="col-tools"><div class="row-actions"><button class="light" data-history="${escapeText(record["记录ID"])}">历史</button><button class="light" data-edit="${escapeText(record["记录ID"])}">编辑</button></div></td>
        </tr>`;
      }).join("");
      $("empty").hidden = rows.length > 0;
      tbody.querySelectorAll("[data-edit]").forEach((btn) => {
        btn.addEventListener("click", () => editRecord(btn.dataset.edit));
      });
      tbody.querySelectorAll("[data-history]").forEach((btn) => {
        btn.addEventListener("click", () => showHistory(btn.dataset.history));
      });
    }
    function chartItemHtml(chart, activeName = "") {
      const linked = chart.inRegistry ? `已入台账 · ${escapeText(chart.status || "-")} · ${escapeText(chart.action || "-")}` : "未入台账";
      const active = chart.name === activeName ? " active" : "";
      return `<div class="chart-item${active}" data-chart-name="${escapeText(chart.name)}">
        <div class="chart-name">${escapeText(chart.name)}</div>
        <div class="chart-meta">账号 ${escapeText(chart.account)} · ${linked}<br>${escapeText(chart.mtime)} · ${escapeText(chart.sizeText || "")}</div>
        <div class="chart-actions">
          <button class="light" data-preview-chart="${escapeText(chart.name)}">预览</button>
          <a class="button-link light" href="${escapeText(chart.url)}" target="_blank" rel="noopener">打开</a>
        </div>
      </div>`;
    }
    async function loadAccountCharts(recordId) {
      $("accountCharts").innerHTML = '<div class="muted">加载中...</div>';
      $("chartPreview").hidden = true;
      $("chartPreview").removeAttribute("src");
      try {
        const data = await requestJson(`/api/accounts/${encodeURIComponent(recordId)}/charts`);
        state.accountCharts = data.charts || [];
        $("chartHint").textContent = state.accountCharts.length ? `${state.accountCharts.length} 个图表` : "暂无图表";
        if (!state.accountCharts.length) {
          $("accountCharts").innerHTML = '<div class="muted">暂无匹配图表，可点击顶部“生成图表”上传 statement。</div>';
          return;
        }
        $("accountCharts").innerHTML = state.accountCharts.map((chart) => chartItemHtml(chart)).join("");
        bindChartButtons($("accountCharts"), state.accountCharts);
      } catch (err) {
        $("accountCharts").innerHTML = `<div class="muted">${escapeText(err.message)}</div>`;
      }
    }
    function bindChartButtons(container, charts) {
      container.querySelectorAll("[data-preview-chart]").forEach((btn) => {
        btn.addEventListener("click", () => {
          const chart = charts.find((item) => item.name === btn.dataset.previewChart);
          if (!chart) return;
          $("chartPreview").src = chart.url;
          $("chartPreview").hidden = false;
          container.querySelectorAll(".chart-item").forEach((item) => item.classList.toggle("active", item.dataset.chartName === chart.name));
        });
      });
    }
    function clearForm() {
      state.editingId = null;
      $("formTitle").textContent = "新增记录";
      Object.values(fields).forEach((id) => $(id).value = "");
      $("type").value = "账户";
      $("status").value = "待复核";
      $("customAction").value = "";
      $("joinedAt").value = "";
      $("updatedAtField").value = "";
      syncCustomAction();
      $("deleteBtn").hidden = true;
      $("historyBtn").hidden = true;
      $("chartHint").textContent = "选择账号后显示";
      $("accountCharts").innerHTML = "";
      $("chartPreview").hidden = true;
      $("chartPreview").removeAttribute("src");
      setStatus("");
    }
    function syncCustomAction() {
      const show = $("action").value === "自定义";
      $("customAction").style.display = show ? "block" : "none";
    }
    function editRecord(id) {
      const record = state.records.find((item) => item["记录ID"] === id);
      if (!record) return;
      state.editingId = id;
      $("formTitle").textContent = record["账号"] ? `编辑 ${record["账号"]}` : "编辑记录";
      for (const [name, idName] of Object.entries(fields)) {
        if (name === "建议动作") ensureSelectOption($(idName), record[name]);
        $(idName).value = record[name] || "";
      }
      $("joinedAt").value = record["加入时间"] || "";
      $("updatedAtField").value = record["修改时间"] || "";
      $("customAction").value = "";
      syncCustomAction();
      $("deleteBtn").hidden = false;
      $("historyBtn").hidden = false;
      loadAccountCharts(id);
      setStatus("");
    }
    function formPayload() {
      const payload = {};
      for (const [name, idName] of Object.entries(fields)) payload[name] = $(idName).value.trim();
      if ($("action").value === "自定义") payload["建议动作"] = $("customAction").value.trim() || "自定义";
      return payload;
    }
    async function requestJson(url, options) {
      const res = await fetch(url, options);
      const data = await res.json();
      if (!res.ok || !data.ok) throw new Error(data.error || `HTTP ${res.status}`);
      return data;
    }
    async function loadData() {
      const data = await requestJson("/api/accounts");
      state.records = data.records || [];
      state.meta = data.summary || {};
      state.historyCounts = data.historyCounts || {};
      renderSummary(state.meta);
      renderRows();
    }
    async function loadMeta() {
      const data = await requestJson("/api/meta");
      fillSelect($("actionFilter"), data.actions || [], "全部建议");
      fillSelect($("statusFilter"), data.statuses || [], "全部状态");
      fillSelect($("action"), data.actions || []);
      fillSelect($("status"), data.statuses || []);
      fillSelect($("type"), data.types || []);
    }
    async function saveCurrent() {
      try {
        const payload = formPayload();
        const url = state.editingId ? `/api/accounts/${encodeURIComponent(state.editingId)}` : "/api/accounts";
        const method = state.editingId ? "PUT" : "POST";
        await requestJson(url, {
          method,
          headers: { "Content-Type": "application/json; charset=utf-8" },
          body: JSON.stringify(payload),
        });
        await loadData();
        clearForm();
        setStatus("已保存");
      } catch (err) {
        setStatus(err.message, false);
      }
    }
    async function deleteCurrent() {
      if (!state.editingId) return;
      const record = state.records.find((item) => item["记录ID"] === state.editingId);
      const label = record ? (record["账号"] || record["关联账号/主体"] || record["记录ID"]) : state.editingId;
      if (!confirm(`删除 ${label}？`)) return;
      try {
        await requestJson(`/api/accounts/${encodeURIComponent(state.editingId)}`, { method: "DELETE" });
        await loadData();
        clearForm();
        setStatus("已删除");
      } catch (err) {
        setStatus(err.message, false);
      }
    }
    function renderChartLibrary() {
      const q = $("chartSearch").value.trim().toLowerCase();
      const filter = $("chartLinkedFilter").value;
      const charts = state.allCharts.filter((chart) => {
        const hay = `${chart.account} ${chart.name}`.toLowerCase();
        return (!q || hay.includes(q)) &&
          (!filter || (filter === "linked" ? chart.inRegistry : !chart.inRegistry));
      });
      $("chartLibraryTitle").textContent = `图表库 (${charts.length}/${state.allCharts.length})`;
      $("chartLibraryList").innerHTML = charts.length
        ? charts.map((chart) => chartItemHtml(chart)).join("")
        : '<div class="empty">没有匹配图表</div>';
      bindChartButtons($("chartLibraryList"), charts);
    }
    async function showChartLibrary() {
      $("chartLibraryList").innerHTML = '<div class="muted">加载中...</div>';
      $("chartLibraryDialog").showModal();
      try {
        const data = await requestJson("/api/charts");
        state.allCharts = data.charts || [];
        renderChartLibrary();
      } catch (err) {
        $("chartLibraryList").innerHTML = `<div class="muted">${escapeText(err.message)}</div>`;
      }
    }
    function parseHistoryJson(text) {
      if (!text) return null;
      try {
        return JSON.parse(text);
      } catch {
        return null;
      }
    }
    function valueOrDash(value) {
      const text = String(value || "").trim();
      return text || "-";
    }
    function versionPanel(title, record, changedFields, tone) {
      if (!record) {
        return `<div class="version-panel ${tone}"><h4>${escapeText(title)}</h4><div class="version-grid"><div class="version-label">状态</div><div class="version-value">无</div></div></div>`;
      }
      const fields = [
        ["账号", record["账号"] || record["关联账号/主体"] || record["记录ID"]],
        ["建议", record["建议动作"]],
        ["状态", record["状态"]],
        ["分组", record["当前分组"]],
        ["标签", record["风险标签"]],
        ["加入时间", record["加入时间"]],
        ["修改时间", record["修改时间"]],
        ["备注", record["风险/问题备注"]],
      ];
      return `<div class="version-panel ${tone}">
        <h4>${escapeText(title)}</h4>
        <div class="version-grid">
          ${fields.map(([label, value]) => {
            const changed = changedFields.has(label) || changedFields.has(label === "建议" ? "建议动作" : label === "分组" ? "当前分组" : label === "标签" ? "风险标签" : label === "备注" ? "风险/问题备注" : label);
            return `<div class="version-label">${escapeText(label)}</div><div class="version-value ${changed ? "changed" : ""}">${escapeText(valueOrDash(value))}</div>`;
          }).join("")}
        </div>
      </div>`;
    }
    function historyItemHtml(item, index) {
      const before = parseHistoryJson(item["修改前JSON"]);
      const after = parseHistoryJson(item["修改后JSON"]);
      const changedFields = new Set(String(item["修改字段"] || "").split(/[；;,，]+/).map((x) => x.trim()).filter(Boolean));
      const operation = item["操作"] || "修改";
      const badgeClass = operation === "删除" ? "version-badge delete" : "version-badge";
      const account = item["账号"] || before?.["账号"] || after?.["账号"] || before?.["关联账号/主体"] || after?.["关联账号/主体"] || item["记录ID"];
      return `<div class="history-item">
        <div class="version-head">
          <div>
            <div class="version-title">
              <span class="${badgeClass}">${escapeText(operation)}</span>
              <b>${escapeText(account || `版本 ${index + 1}`)}</b>
            </div>
            <div class="time-text">修改字段：${escapeText(item["修改字段"] || "-")}</div>
          </div>
          <div class="version-time">操作时间<br>${escapeText(item["修改时间"] || "-")}</div>
        </div>
        <div class="version-body">
          ${versionPanel(operation === "删除" ? "删除前版本" : "修改前版本", before, changedFields, "before")}
          ${versionPanel(operation === "删除" ? "删除后版本" : "修改后版本", after, changedFields, "after")}
        </div>
      </div>`;
    }
    async function showHistory(recordId) {
      try {
        const record = state.records.find((item) => item["记录ID"] === recordId);
        const data = await requestJson(`/api/accounts/${encodeURIComponent(recordId)}/history`);
        $("historyTitle").textContent = `修改历史 ${record ? (record["账号"] || record["关联账号/主体"] || "") : ""}`;
        if (!data.history.length) {
          $("historyBody").innerHTML = '<div class="empty">暂无修改历史</div>';
        } else {
          $("historyBody").innerHTML = data.history.map((item, index) => historyItemHtml(item, index)).join("");
        }
        $("historyDialog").showModal();
      } catch (err) {
        setStatus(err.message, false);
      }
    }
    function resetPageAndRender() {
      state.page = 1;
      renderRows();
    }
    $("search").addEventListener("input", resetPageAndRender);
    $("actionFilter").addEventListener("change", resetPageAndRender);
    $("statusFilter").addEventListener("change", resetPageAndRender);
    $("sortBy").addEventListener("change", resetPageAndRender);
    $("prevPageBtn").addEventListener("click", () => {
      if (state.page > 1) {
        state.page -= 1;
        renderRows();
      }
    });
    $("nextPageBtn").addEventListener("click", () => {
      const pages = Math.max(1, Math.ceil(state.lastFilteredCount / state.pageSize));
      if (state.page < pages) {
        state.page += 1;
        renderRows();
      }
    });
    $("action").addEventListener("change", syncCustomAction);
    $("newBtn").addEventListener("click", clearForm);
    $("reloadBtn").addEventListener("click", loadData);
    $("clearBtn").addEventListener("click", clearForm);
    $("saveBtn").addEventListener("click", saveCurrent);
    $("deleteBtn").addEventListener("click", deleteCurrent);
    $("historyBtn").addEventListener("click", () => state.editingId && showHistory(state.editingId));
    $("closeHistoryBtn").addEventListener("click", () => $("historyDialog").close());
    $("chartLibraryBtn").addEventListener("click", showChartLibrary);
    $("closeChartLibraryBtn").addEventListener("click", () => $("chartLibraryDialog").close());
    $("chartSearch").addEventListener("input", renderChartLibrary);
    $("chartLinkedFilter").addEventListener("change", renderChartLibrary);
    (async function init() {
      try {
        await loadMeta();
        await loadData();
        clearForm();
      } catch (err) {
        setStatus(err.message, false);
      }
    })();
  </script>
</body>
</html>
"""


def main() -> None:
    init_workbook()
    normalize_initial_record_times()
    server = ThreadingHTTPServer((HOST, PORT), Handler)
    print(f"问题账户台账网页已启动: http://{HOST}:{PORT}")
    print(f"Excel: {WORKBOOK_PATH}")
    server.serve_forever()


if __name__ == "__main__":
    main()
